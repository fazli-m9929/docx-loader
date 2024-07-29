from docx.oxml.text.paragraph import CT_P
from lxml.etree import ElementTree, _Element
from docxlatex import Document as DC 
from docx import Document
from docx.table import Table
from io import StringIO
import langchain_core.documents
import re
import os

PARAGRAPH_TAG = 'p'
TABLE_TAG = 'tbl'
SECTPR_TAG = 'sectPr'
SDT_TAG = 'sdt'

class DocumentLoader:
    def __init__(self, doc_path):
        self.doc_path = doc_path
        self.document_obj = Document(doc_path)
    
    def contains_mathml(self, element: CT_P):
        xml_str = element.xml
        return '<m:' in xml_str

    def xml_to_text(self, element: CT_P):
        xml_str = element.xml
        xml_to_text = DC("").xml_to_text
        latex_text = xml_to_text(xml_str).replace('\n', '')

        pattern = r'\$(.*?)\$'
        def remove_spaces(match):
            return f"${match.group(1).replace(' ', '')}$"
        
        return re.sub(pattern, remove_spaces, latex_text)

    def extract_toc_entries(self, xml_tree: _Element):
        root = ElementTree(xml_tree).getroot()
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        }

        toc_entries = []

        for para in root.findall('.//w:p', namespaces):
            para_style = para.find('w:pPr/w:pStyle', namespaces)
            if para_style is not None and para_style.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '').startswith('TOC'):
                text_elements = para.findall('.//w:t', namespaces)
                toc_text = '\t\tPage '.join([t.text or '' for t in text_elements if t.text])
                if toc_text:
                    toc_entries.append(toc_text)

        return '\n'.join(toc_entries)

    def table_to_plain_text(self, table: Table):
        output = StringIO()
        processed_cells = set()

        previous_cell_id = (0, 0)
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_id = cell._element.getparent().getparent().index(cell._element.getparent()), cell._element.getparent().index(cell._element)
                if cell_id in processed_cells:
                    if cell_id == previous_cell_id:
                        row_text.append("merged_row")
                    else:
                        row_text.append('\t')
                else:
                    if self.contains_mathml(cell._element):
                        cell.text = self.xml_to_text(cell._element)
                    row_text.append(cell.text)
                    processed_cells.add(cell_id)
                previous_cell_id = cell_id
            if all(cell_text == '' for cell_text in row_text):
                continue
            row_text = [item for item in row_text if item !='merged_row']
            output.write("[ " +" | ".join(row_text) + " ]" + "\n")

        def remove_extra_newlines(text: str):
            cleaned_text = re.sub(r'\n{2,}', '\n', text)
            return cleaned_text

        return remove_extra_newlines(output.getvalue())

    def create_element_index_dict(self):
        body_element = self.document_obj._body._element

        para_idx = 0
        tbl_idx = 0
        
        index_dict = {}

        for idx, elem in enumerate(body_element):
            tag = elem.tag.split("}")[-1]

            if tag == PARAGRAPH_TAG:
                index_dict[idx] = (tag, para_idx)
                para_idx += 1
            elif tag == TABLE_TAG:
                index_dict[idx] = (tag, tbl_idx)
                tbl_idx += 1
            else:
                index_dict[idx] = (tag, None)

        return index_dict

    def generate_text_list(self, index_dict, tag_flag=True):
        text_list = []

        for key, (tag, index) in index_dict.items():
            if tag == SDT_TAG:
                text = self.extract_toc_entries(self.document_obj._body._element[key])

            elif tag == PARAGRAPH_TAG:
                if self.contains_mathml(self.document_obj._body._element[key]):
                    self.document_obj.paragraphs[index].text = self.xml_to_text(self.document_obj._body._element[key])
                text = self.document_obj.paragraphs[index].text

            elif tag == TABLE_TAG:
                text = self.table_to_plain_text(self.document_obj.tables[index])

            else:
                text = self.document_obj._body._element[key].text

            if text is None:
                continue
            if tag_flag:
                text_list.append((text, tag))
            else:
                text_list.append(text)

        return text_list

    def load(self):
        index_dict = self.create_element_index_dict()
        text_list = self.generate_text_list(index_dict, False)
        document = langchain_core.documents.Document(
            page_content=re.sub(r'\n{3,}', '\n\n', '\n'.join(text_list)),
            metadata={'source': os.path.basename(self.doc_path)}
        )
        return [document]