from docx.oxml.text.paragraph import CT_P
from lxml.etree import ElementTree, _Element
from docxlatex import Document
from docx.table import Table
from io import StringIO
import docx.document

PARAGRAPH_TAG = 'p'
TABLE_TAG = 'tbl'
SECTPR_TAG = 'sectPr'
SDT_TAG = 'sdt'


def contains_mathml(element: CT_P):
    xml_str = element.xml
    return '<m:' in xml_str


def xml_to_text(element: CT_P):
    xml_str = element.xml
    xml_to_text = Document("").xml_to_text
    latex_text = xml_to_text(xml_str)
    return latex_text.replace('\n', '').replace(' ','')


def extract_toc_entries(xml_tree: _Element):
    root = ElementTree(xml_tree).getroot()
    namespaces = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    }

    toc_entries = []

    # Iterate over all paragraph elements
    for para in root.findall('.//w:p', namespaces):
        # Find the paragraph style
        para_style = para.find('w:pPr/w:pStyle', namespaces)
        if para_style is not None and para_style.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '').startswith('TOC'):
            text_elements = para.findall('.//w:t', namespaces)
            toc_text = '\t\tPage '.join([t.text or '' for t in text_elements if t.text])
            if toc_text:  # Avoid adding empty entries
                toc_entries.append(toc_text)

    return '\n'.join(toc_entries)


def table_to_plain_text(table: Table):
    output = StringIO()
    processed_cells = set()

    # Process the cells to handle MathML and convert them to plain text
    previous_cell_id = (0, 0)
    for row in table.rows:
        row_text = []
        for cell in row.cells:
            cell_id = cell._element.getparent().getparent().index(cell._element.getparent()), cell._element.getparent().index(cell._element)
            if cell_id in processed_cells:
                if cell_id == previous_cell_id:
                    row_text.append("merged_row")
                else:
                    row_text.append('\t')
            else:
                if contains_mathml(cell._element):  # Assuming this function exists
                    cell.text = xml_to_text(cell._element)  # Assuming this function exists
                row_text.append(cell.text)
                processed_cells.add(cell_id)
            previous_cell_id = cell_id
        # Skip the row if all elements are empty strings
        if all(cell_text == '' for cell_text in row_text):
            continue
        # Join the cell texts with commas and add a newline at the end
        row_text = [item for item in row_text if item !='merged_row']
        output.write("[ " +" | ".join(row_text) + " ]" + "\n")

    return output.getvalue()


def create_element_index_dict(doc: docx.document.Document):
    body_element = doc._body._element

    PARAGRAPH_TAG = 'p'
    TABLE_TAG = 'tbl'
    
    para_idx = 0
    tbl_idx = 0
    
    index_dict = {}

    for idx, elem in enumerate(body_element):
        tag = elem.tag.split("}")[-1]  # Extract tag name (ignore namespace)

        if tag == PARAGRAPH_TAG:
            index_dict[idx] = (tag, para_idx)
            para_idx += 1
        elif tag == TABLE_TAG:
            index_dict[idx] = (tag, tbl_idx)
            tbl_idx += 1
        else:
            index_dict[idx] = (tag, None)

    return index_dict


def generate_text_list(index_dict, document_obj):
    text_list = []

    for key, (tag, index) in index_dict.items():
        if tag == SDT_TAG:
            text = extract_toc_entries(document_obj._body._element[key])

        elif tag == PARAGRAPH_TAG:
            if contains_mathml(document_obj._body._element[key]):
                document_obj.paragraphs[index].text = xml_to_text(document_obj._body._element[key])
            text = document_obj.paragraphs[index].text

        elif tag == TABLE_TAG:
            text = table_to_plain_text(document_obj.tables[index])

        else:
            text = document_obj._body._element[key].text

        if text is None:
            continue
        text_list.append((text, tag))

    return text_list


def combine_tables_with_captions(text_list):
    combined_list = []
    i = 0
    while i < len(text_list):
        if text_list[i][1] == 'tbl':
            # If the list starts with a table, don't combine it
            combined_list.append(text_list[i])
            i += 1
        elif text_list[i][1] == 'p' and (text_list[i][0].startswith('table') or text_list[i][0].startswith('جدول')):
            # Check if the 'p' tagged text starts with 'table' or 'جدول'
            combined_item = text_list[i]
            if i + 1 < len(text_list) and text_list[i + 1][1] == 'tbl':
                combined_item = (combined_item[0] + '\n' + text_list[i + 1][0], 'tbl')
                i += 2
            elif i + 2 < len(text_list) and text_list[i + 2][1] == 'tbl':
                combined_item = (combined_item[0] + '\n' + text_list[i + 1][0] + ' ' + text_list[i + 2][0], 'tbl')
                i += 3
            combined_list.append(combined_item)
        else:
            # Combine table with the 'p' before it
            if i + 1 < len(text_list) and text_list[i + 1][1] == 'tbl':
                combined_item = (text_list[i][0] + '\n' + text_list[i + 1][0], 'tbl')
                combined_list.append(combined_item)
                i += 2
            else:
                combined_list.append(text_list[i])
                i += 1
    return combined_list


